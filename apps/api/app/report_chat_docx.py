"""Native Word rendering for validated persona report documents."""

from __future__ import annotations

from io import BytesIO
import re

from docx import Document

from .report_chat_schemas import ReportDocument


def safe_docx_filename(title: str) -> str:
    stem = re.sub(r"[^A-Za-z0-9._ -]+", "", title).strip(" .")
    return f"{stem or 'report'}.docx"


class ReportChatDocxRenderer:
    def render(self, report: ReportDocument) -> bytes:
        document = Document()
        document.add_heading(report.title, level=0)
        for block in report.blocks:
            suffix = (
                " " + " ".join(f"[{item}]" for item in block.citation_ids)
                if block.citation_ids
                else ""
            )
            if block.kind == "heading":
                document.add_heading(block.text, level=block.level)
            elif block.kind == "paragraph":
                document.add_paragraph(block.text + suffix)
            elif block.kind in {"bullet_list", "numbered_list"}:
                style = (
                    "List Bullet"
                    if block.kind == "bullet_list"
                    else "List Number"
                )
                for item in block.items:
                    document.add_paragraph(item + suffix, style=style)
            elif block.kind == "table":
                table = document.add_table(
                    rows=1,
                    cols=len(block.columns),
                )
                for index, label in enumerate(block.columns):
                    table.rows[0].cells[index].text = label
                for values in block.rows:
                    cells = table.add_row().cells
                    for index, value in enumerate(values):
                        cells[index].text = value
                document.add_paragraph(suffix.strip())

        document.add_heading("Evidence Sources", level=1)
        for citation in report.citations:
            document.add_paragraph(
                f"[{citation.id}] {citation.label} — {citation.source_ref}"
            )

        buffer = BytesIO()
        document.save(buffer)
        return buffer.getvalue()
